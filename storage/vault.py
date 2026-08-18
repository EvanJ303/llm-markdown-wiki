import json
from dataclasses import dataclass
import sqlite3
from pathlib import Path
from typing import List

@dataclass
class Document:
    path: Path
    type: str
    size: int
    created: str
    modified: str
    hash: str
    processed: bool

@dataclass
class Chunk:
    content: str
    page: int | None = None

@dataclass
class SearchHit:
    path: Path
    snippet: str
    page: int | None = None

@dataclass
class DocumentContent:
    content: str
    created: str
    modified: str
    path: Path

class Vault:
    def __init__(self, root: Path) -> None:
        self.root = root.resolve()

        self.wiki_path = self.root / 'wiki'
        self.db_path = self.root / '.llmwiki/index.db'

        self.wiki_path.mkdir(parents=True, exist_ok=True)

        config_path = Path(__file__).resolve().parent.parent / 'config.json'
        with open(config_path, 'r', encoding='utf-8') as f:
            cfg = json.load(f)
            self.chunk_size = cfg.get('chunk_size', 200)
            self.chunk_overlap = cfg.get('chunk_overlap', 20)
            self.min_chunk_tokens = cfg.get('min_chunk_tokens', 50)
            self.max_chunk_chars = cfg.get('max_chunk_chars', 2000)

        self.conn = sqlite3.connect(self.db_path)
        
        schema_path = Path(__file__).resolve().parent / 'schema.sql'
        with open(schema_path, 'r', encoding='utf-8') as f:
            schema = f.read()

        self.conn.executescript(schema)

    def _extract_text(self, path: Path) -> List[tuple[str, int | None]]:
        suffix = path.suffix.lower()
        try:
            if suffix == '.pdf':
                from pypdf import PdfReader

                reader = PdfReader(str(path))
                parts: List[tuple[str, int | None]] = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        parts.append((text, i + 1))
                return parts

            if suffix == '.docx':
                from docx import Document as DocxDocument

                doc = DocxDocument(path)
                # docx has no reliable concept of page; return as single segment
                text = '\n'.join(p.text for p in doc.paragraphs if p.text)
                return [(text, None)]

            if suffix in ('.xlsx', '.xlsm', '.xltx', '.xltm'):
                from openpyxl import load_workbook

                wb = load_workbook(filename=path, read_only=True, data_only=True)
                parts_ret: List[tuple[str, int | None]] = []
                for i, sheet in enumerate(wb.worksheets):
                    lines: List[str] = []
                    lines.append(f'=== Sheet: {sheet.title} ===')
                    for row in sheet.iter_rows(values_only=True):
                        row_vals = [str(cell) if cell is not None else '' for cell in row]
                        lines.append('\t'.join(row_vals))
                    parts_ret.append(('\n'.join(lines), i + 1))
                return parts_ret

            # fallback: try to read as plain text and detect form-feed pages
            try:
                raw = path.read_text(encoding='utf-8', errors='ignore')
                if '\f' in raw:
                    pages = [p for p in raw.split('\f') if p.strip()]
                    return [(p, i + 1) for i, p in enumerate(pages)]
                return [(raw, None)]
            except Exception:
                return []
        except Exception:
            return []

    def _read_and_chunk_text(self, path: Path) -> list[Chunk]:
        segments = self._extract_text(path)

        def est_tokens(s: str) -> int:
            return max(1, len(s) // 4)

        def split_long_text(s: str, page: int | None) -> List[Chunk]:
            """Split a very long text block into smaller chunks by char windows with overlap."""
            res: List[Chunk] = []
            maxc = int(self.max_chunk_chars)
            overlap_chars = int(self.chunk_overlap * 4)
            if maxc <= 0:
                maxc = 2000
            start = 0
            length = len(s)
            while start < length:
                end = min(start + maxc, length)
                piece = s[start:end].strip()
                if piece:
                    res.append(Chunk(content=piece, page=page))
                if end >= length:
                    break
                # move start to create overlap
                start = max(0, end - overlap_chars)
            return res

        chunks: List[Chunk] = []

        for (segment_text, page_num) in segments:
            if not segment_text or not segment_text.strip():
                continue

            # split into paragraphs
            import re

            # If this segment originates from a spreadsheet (we add a header '=== Sheet:'),
            # treat each row/line as a paragraph so chunks follow row boundaries.
            if segment_text.lstrip().startswith('=== Sheet:'):
                paragraphs = [l.strip() for l in segment_text.splitlines() if l.strip()]
            else:
                paragraphs = [p.strip() for p in re.split(r'\n\s*\n', segment_text) if p.strip()]

            current = ''

            for paragraph in paragraphs:
                # split paragraph into sentences (simple regex)
                sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', paragraph) if s.strip()]

                # if paragraph is a single very long sentence, treat it as a long block
                if len(sentences) == 1 and (len(sentences[0]) > self.max_chunk_chars or est_tokens(sentences[0]) > self.chunk_size):
                    # flush current if present
                    if current:
                        chunks.append(Chunk(content=current.strip(), page=page_num))
                        current = ''
                    chunks.extend(split_long_text(sentences[0], page_num))
                    continue

                for sent in sentences:
                    if not current:
                        current = sent
                        continue

                    tentative = current + ' ' + sent
                    if (len(tentative) <= self.max_chunk_chars and est_tokens(tentative) <= self.chunk_size) or est_tokens(current) < self.min_chunk_tokens:
                        current = tentative
                        continue

                    # finalize current chunk
                    chunk_text = current.strip()
                    if chunk_text:
                        chunks.append(Chunk(content=chunk_text, page=page_num))

                    # prepare overlap for next chunk
                    overlap_chars = int(self.chunk_overlap * 4)
                    if overlap_chars > 0:
                        overlap_text = chunk_text[-overlap_chars:]
                        # start next with overlap + current sentence
                        current = (overlap_text + ' ' + sent).strip()
                    else:
                        current = sent

                # paragraph finished; if it's long relative to limits, flush
                if current and (est_tokens(current) >= self.chunk_size or len(current) >= self.max_chunk_chars):
                    chunks.append(Chunk(content=current.strip(), page=page_num))
                    current = ''

            # end of page/segment: flush remaining current
            if current:
                # if still too long, split further
                if len(current) > self.max_chunk_chars or est_tokens(current) > self.chunk_size:
                    chunks.extend(split_long_text(current, page_num))
                else:
                    chunks.append(Chunk(content=current.strip(), page=page_num))

        # ensure every chunk has content and trim
        final_chunks: List[Chunk] = []
        for c in chunks:
            content = (c.content or '').strip()
            if not content:
                continue
            # truncate to max chars if needed
            if len(content) > self.max_chunk_chars:
                # split into safe pieces
                final_chunks.extend(split_long_text(content, c.page))
            else:
                final_chunks.append(Chunk(content=content, page=c.page))

        return final_chunks

    def _upsert_document_row(self, document: Document) -> int:
        self.conn.execute(
            '''
            INSERT INTO documents (path, type, size, created, modified, hash, processed)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(path) DO UPDATE SET
                type = excluded.type,
                size = excluded.size,
                created = excluded.created,
                modified = excluded.modified,
                hash = excluded.hash,
                processed = excluded.processed
            ''',
            (
                str(document.path),
                document.type,
                document.size,
                document.created,
                document.modified,
                document.hash,
                document.processed,
            ),
        )
        self.conn.commit()

        row = self.conn.execute(
            'SELECT id FROM documents WHERE path = ?',
            (str(document.path),),
        ).fetchone()
        
        return int(row[0])

    def _insert_chunk_row(self, chunk: Chunk) -> None:
        self.conn.execute(
            '''
            INSERT INTO chunks (content, page)
            VALUES (?, ?)
            ''',
            (
                chunk.content,
                chunk.page,
            ),
        )
        self.conn.commit()

    def _is_document_unchanged(self, path: Path) -> bool:
        try:
            import hashlib

            p = Path(path).resolve()
            stat = p.stat()
            current_mtime = str(stat.st_mtime)

            row = self.conn.execute(
                'SELECT modified, hash FROM documents WHERE path = ?',
                (str(p),),
            ).fetchone()
            if row is None:
                return False

            stored_mtime, stored_hash = row
            if stored_mtime == current_mtime:
                return True

            current_hash = hashlib.sha256(p.read_bytes()).hexdigest()
            return stored_hash == current_hash
        except Exception:
            return False

    def _index_document(self, path: Path) -> None:
        try:
            import hashlib

            p = Path(path).resolve()

            if self._is_document_unchanged(p):
                return

            try:
                stat = p.stat()
            except Exception:
                stat = None

            chunks = []
            try:
                chunks = self._read_and_chunk_text(p)
            except Exception:
                chunks = []

            doc = Document(
                path=p,
                type=p.suffix.lower().lstrip('.'),
                size=(stat.st_size if stat is not None else 0),
                created=(str(stat.st_birthtime) if stat is not None else ''),
                modified=(str(stat.st_mtime) if stat is not None else ''),
                hash=(hashlib.sha256(p.read_bytes()).hexdigest() if stat is not None else ''),
                processed=True,
            )

            doc_id = self._upsert_document_row(doc)

            try:
                self.conn.execute('DELETE FROM chunks WHERE document_id = ?', (doc_id,))
            except Exception:
                pass

            for chunk in chunks:
                try:
                    if not chunk.content or not str(chunk.content).strip():
                        continue
                    self.conn.execute(
                        'INSERT INTO chunks (document_id, content, page) VALUES (?, ?, ?)',
                        (doc_id, chunk.content, chunk.page),
                    )
                except Exception:
                    continue

            self.conn.commit()

            try:
                self.conn.execute("INSERT INTO chunks_fts(chunks_fts) VALUES('rebuild')")
                self.conn.commit()
            except Exception:
                pass

        except Exception:
            return

    def _remove_document(self, path: Path) -> None:
        try:
            p = Path(path).resolve()
            try:
                self.conn.execute('DELETE FROM documents WHERE path = ?', (str(p),))
                self.conn.commit()
            except Exception:
                pass
        except Exception:
            return

    def index_db(self) -> None:
        try:
            for p in self.wiki_path.rglob('*'):
                if not p.is_file():
                    continue
                try:
                    if self._is_document_unchanged(p):
                        continue
                    self._index_document(p)
                except Exception:
                    continue

            try:
                rows = self.conn.execute('SELECT path FROM documents').fetchall()
                for (doc_path,) in rows:
                    try:
                        if not Path(doc_path).exists():
                            self._remove_document(Path(doc_path))
                    except Exception:
                        continue
            except Exception:
                pass

        except Exception:
            return

    def search(self, query: str) -> list[SearchHit]:
        hits: List[SearchHit] = []
        if not query or not query.strip():
            return hits

        q = query.strip()

        try:
            # Try full-text search with FTS5 and produce a highlighted snippet
            sql = (
                "SELECT documents.path as path, "
                "snippet(chunks_fts, '<b>', '</b>', '...', 64) as snippet, "
                "chunks.page as page "
                "FROM chunks_fts "
                "JOIN chunks ON chunks_fts.rowid = chunks.chunk_id "
                "JOIN documents ON chunks.document_id = documents.id "
                "WHERE chunks_fts MATCH ? "
                "LIMIT 200"
            )
            cur = self.conn.execute(sql, (q,))
            for row in cur.fetchall():
                path_str, snippet_text, page = row
                if not snippet_text:
                    snippet_text = ''
                hits.append(SearchHit(path=Path(path_str), snippet=snippet_text, page=(page if page is not None else None)))
            return hits
        except Exception:
            # Fallback to simple LIKE search if FTS not available
            try:
                like_q = f"%{q}%"
                cur = self.conn.execute(
                    'SELECT documents.path, chunks.content, chunks.page FROM chunks JOIN documents ON chunks.document_id = documents.id WHERE chunks.content LIKE ? LIMIT 200',
                    (like_q,)
                )
                for path_str, content, page in cur.fetchall():
                    snippet_text = ''
                    try:
                        lower = content.lower()
                        idx = lower.find(q.lower())
                        if idx >= 0:
                            start = max(0, idx - 64)
                            end = min(len(content), idx + len(q) + 64)
                            prefix = '...' if start > 0 else ''
                            suffix = '...' if end < len(content) else ''
                            snippet_text = f"{prefix}{content[start:end]}{suffix}"
                        else:
                            snippet_text = (content[:128] + '...') if len(content) > 128 else content
                    except Exception:
                        snippet_text = (content[:128] + '...') if len(content) > 128 else content

                    hits.append(SearchHit(path=Path(path_str), snippet=snippet_text, page=(page if page is not None else None)))
                return hits
            except Exception:
                return hits

    def read_document(self, path: Path, page_start: int | None = None, page_end: int | None = None) -> DocumentContent:
        p = Path(path).resolve()

        doc_row = self.conn.execute(
            'SELECT created, modified FROM documents WHERE path = ?',
            (str(p),),
        ).fetchone()

        if doc_row is None:
            created = str(p.stat().st_ctime) if p.exists() else ''
            modified = str(p.stat().st_mtime) if p.exists() else ''
        else:
            created, modified = doc_row

        query = 'SELECT content, page FROM chunks JOIN documents ON chunks.document_id = documents.id WHERE documents.path = ?'
        params: List[object] = [str(p)]

        if page_start is not None or page_end is not None:
            clauses = []
            if page_start is not None:
                clauses.append('chunks.page >= ?')
                params.append(page_start)
            if page_end is not None:
                clauses.append('chunks.page <= ?')
                params.append(page_end)
            query += ' AND ' + ' AND '.join(clauses)

        query += ' ORDER BY chunks.page ASC, chunks.rowid ASC'

        rows = self.conn.execute(query, params).fetchall()
        if not rows:
            # fallback: if database has no rows yet, read the source file directly
            text = self._extract_text(p)
            raw_parts: List[str] = []
            for segment, page in text:
                if page_start is not None and page is not None and page < page_start:
                    continue
                if page_end is not None and page is not None and page > page_end:
                    continue
                raw_parts.append(segment)
            content = '\n\n'.join(raw_parts).strip()
            return DocumentContent(content=content, created=created, modified=modified, path=p)

        content_parts: List[str] = []
        for chunk_content, page in rows:
            if page_start is not None and page is not None and page < page_start:
                continue
            if page_end is not None and page is not None and page > page_end:
                continue
            if chunk_content:
                content_parts.append(str(chunk_content))

        content = '\n\n'.join(content_parts).strip()
        return DocumentContent(content=content, created=created, modified=modified, path=p)

    def _resolve_wiki_path(self, path: Path) -> Path:
        if path.is_absolute():
            target = path.resolve()
        else:
            target = (self.wiki_path / path).resolve()

        wiki_root = self.wiki_path.resolve()
        try:
            target.relative_to(wiki_root)
        except ValueError as exc:
            raise ValueError(f'page path must be inside the wiki directory: {path}') from exc

        return target

    def _ensure_markdown_path(self, path: Path) -> Path:
        target = Path(path)
        if target.suffix.lower() not in {'.md', '.markdown', '.csv', '.json', '.svg', '.txt'}:
            return target.with_suffix('.md')
        return target

    def _write_page_content(self, target: Path, content: str) -> None:
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding='utf-8')
        self._index_document(target)

    def write_page(self, path: Path, content: str) -> None:
        if not isinstance(content, str):
            raise TypeError('content must be a string')

        target = self._resolve_wiki_path(path)
        target = self._ensure_markdown_path(target)
        self._write_page_content(target, content)

    def edit_page(self, path: Path, pattern: str, replacement: str) -> None:
        if not isinstance(pattern, str):
            raise TypeError('pattern must be a string')
        if not isinstance(replacement, str):
            raise TypeError('replacement must be a string')

        target = self._resolve_wiki_path(path)
        target = self._ensure_markdown_path(target)

        if not target.exists():
            raise FileNotFoundError(f'page does not exist: {target}')

        current_content = target.read_text(encoding='utf-8')
        updated_content = current_content.replace(pattern, replacement)
        self._write_page_content(target, updated_content)

    def delete_page(self, path: Path) -> None:
        target = self._resolve_wiki_path(path)
        target = self._ensure_markdown_path(target)

        if target.exists():
            target.unlink()

        self._remove_document(target)

    def close(self) -> None:
        self.conn.close()